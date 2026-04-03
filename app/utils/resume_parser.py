"""
Resume Parser - Extract text from PDF and DOCX files
"""

import PyPDF2
import docx
import io


def extract_text_from_pdf(uploaded_file):
    """
    Extract text from uploaded PDF file
    
    Parameters:
    uploaded_file : streamlit UploadedFile
        PDF file uploaded via Streamlit
    
    Returns:
    str : Extracted text
    """
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # Reset file pointer for potential re-use
        uploaded_file.seek(0)
        
        if len(text.strip()) == 0:
            return None
        
        return text.strip()
    
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return None


def extract_text_from_docx(uploaded_file):
    """
    Extract text from uploaded DOCX file
    
    Parameters:
    uploaded_file : streamlit UploadedFile
        DOCX file uploaded via Streamlit
    
    Returns:
    str : Extracted text
    """
    try:
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        # Reset file pointer
        uploaded_file.seek(0)
        
        if len(text.strip()) == 0:
            return None
        
        return text.strip()
    
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return None


def extract_text_from_txt(uploaded_file):
    """
    Extract text from uploaded TXT file
    
    Parameters:
    uploaded_file : streamlit UploadedFile
        TXT file uploaded via Streamlit
    
    Returns:
    str : Extracted text
    """
    try:
        text = uploaded_file.read().decode('utf-8')
        uploaded_file.seek(0)
        
        if len(text.strip()) == 0:
            return None
        
        return text.strip()
    
    except Exception as e:
        print(f"Error extracting TXT: {e}")
        return None


def parse_resume(uploaded_file):
    """
    Parse resume from uploaded file (PDF, DOCX, or TXT)
    
    Parameters:
    uploaded_file : streamlit UploadedFile
        File uploaded via Streamlit
    
    Returns:
    tuple : (extracted_text, file_type, error_message)
    """
    if uploaded_file is None:
        return None, None, "No file uploaded"
    
    file_name = uploaded_file.name.lower()
    
    if file_name.endswith('.pdf'):
        text = extract_text_from_pdf(uploaded_file)
        file_type = "PDF"
    elif file_name.endswith('.docx'):
        text = extract_text_from_docx(uploaded_file)
        file_type = "DOCX"
    elif file_name.endswith('.txt'):
        text = extract_text_from_txt(uploaded_file)
        file_type = "TXT"
    else:
        return None, None, "Unsupported file format. Please upload PDF, DOCX, or TXT."
    
    if text is None:
        return None, file_type, f"Could not extract text from {file_type} file. File may be empty or corrupted."
    
    return text, file_type, None


if __name__ == "__main__":
    print("Resume Parser module loaded successfully!")
    print("Supported formats: PDF, DOCX, TXT")